import docx
from docx.api import Document
import docx.shared
from typing import List
import shutil
import re
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
import copy
from lxml import etree
from .schemas import SegmentInternal, TagModel

def _inject_comments(doc: Document, segments: List[SegmentInternal]):
    """
    Updates the comments.xml part of the document with translated text.
    """
    # 1. Map segments by comment_id
    comment_segs = {s.metadata["comment_id"]: s for s in segments if s.metadata.get("type") == "comment"}
    
    if not comment_segs:
        return

    try:
        part = doc.part
        comments_part = None
        for rel in part.rels.values():
             if "comments" in rel.reltype and not "commentsExtended" in rel.reltype and not "commentsIds" in rel.reltype:
                 comments_part = rel.target_part
                 break
                 
        if not comments_part:
            print("Warning: No comments part found to update.")
            return
            
        xml_data = comments_part.blob
        root = etree.fromstring(xml_data)
        namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        updated_count = 0
        
        for comment in root.findall('.//w:comment', namespaces):
             cid = comment.get(qn('w:id'))
             if cid in comment_segs:
                 seg = comment_segs[cid]
                 target_text = seg.target_content if seg.target_content is not None else seg.source_text
                 
                 # Replace content
                 # Clear existing paragraphs/runs in XML
                 # Usually comments have w:p -> w:r -> w:t
                 # We can just replace the text of the first run and delete others?
                 # Or better: clear all children, create new w:p/w:r/w:t?
                 # Since this is low-level XML, let's keep it simple:
                 # Find all w:t and replace valid text?
                 # But target might be shorter/longer or have formatting?
                 # For MVP: Flatten text.
                 
                 # Remove existing content children (p)
                 for child in list(comment):
                     if child.tag == qn('w:p'):
                        comment.remove(child)
                        
                 # Create new simple paragraph
                 # <w:p><w:r><w:t>TEXT</w:t></w:r></w:p>
                 wp = etree.SubElement(comment, qn('w:p'))
                 wr = etree.SubElement(wp, qn('w:r'))
                 wt = etree.SubElement(wr, qn('w:t'))
                 wt.text = target_text
                 
                 updated_count += 1
                 
        if updated_count > 0:
            # Save back
            new_xml = etree.tostring(root, encoding='utf-8', standalone=True)
            comments_part._blob = new_xml
            print(f"Updated {updated_count} comments in comments.xml")
            
    except Exception as e:
        print(f"Error updating comments: {e}")

def _inject_footnotes(doc: Document, segments: List[SegmentInternal]):
    """
    Updates the footnotes.xml part of the document with translated text.
    Uses _inject_tagged_text to preserve formatting/links.
    """
    footnote_segs = {s.metadata["footnote_id"]: s for s in segments if s.metadata.get("type") == "footnote"}
    
    if not footnote_segs:
        return

    try:
        part = doc.part
        footnotes_part = None
        for rel in part.rels.values():
             if "footnotes" in rel.reltype:
                 footnotes_part = rel.target_part
                 break
                 
        if not footnotes_part:
            print("Warning: No footnotes part found to update.")
            return
            
        xml_data = footnotes_part.blob
        root = etree.fromstring(xml_data)
        namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        updated_count = 0
        
        for footnote in root.findall('.//w:footnote', namespaces):
             fid = footnote.get(qn('w:id'))
             if fid in footnote_segs:
                 seg = footnote_segs[fid]
                 target_text = seg.target_content if seg.target_content is not None else seg.source_text
                 
                 # 1. Clear existing content (paragraphs)
                 for child in list(footnote):
                     if child.tag == qn('w:p'):
                        footnote.remove(child)
                 
                 # 2. Create new Paragraph
                 # We need to wrap it in a python-docx Paragraph object to use our helper
                 # But we are operating on lxml elements here.
                 # _inject_tagged_text expects a docx.text.paragraph.Paragraph!
                 
                 # Trick: We can create a dummy docx Paragraph wrapper around our new element?
                 # element = etree.SubElement(footnote, qn('w:p'))
                 # para = docx.text.paragraph.Paragraph(element, part) 
                 # We need 'part' (footnotes_part) which is a Part object.
                 
                 wp = etree.SubElement(footnote, qn('w:p'))
                 # Set style? FootnoteText is standard.
                 pPr = etree.SubElement(wp, qn('w:pPr'))
                 pStyle = etree.SubElement(pPr, qn('w:pStyle'))
                 pStyle.set(qn('w:val'), 'FootnoteText')

                 # Wrap in docx Object
                 from docx.text.paragraph import Paragraph
                 # footnotes_part is a regular Part, but Paragraph expects a parent? 
                 # actually Paragraph(element, parent)
                 # parent is usually the defined parent object (Body, Cell etc).
                 # We can pass None or a mock if _inject_tagged_text only uses .add_run / .clear_content
                 # _inject_tagged_text uses: paragraph._element, paragraph.add_run()
                 # .add_run() needs self._parent to be valid?
                 # Let's check docx source code mental model...
                 # add_run creates Run(r, self). It appends r to p element.
                 # It doesn't seem to strictly require parent for basic operations.
                 
                 # Let's try constructing it with the part as parent (standard behavior usually)
                 proxy_para = Paragraph(wp, footnotes_part)
                 
                 # 3. Inject Content
                 _inject_tagged_text(proxy_para, target_text, seg.tags)
                 
                 # 4. Prepend Footnote Reference
                 # <w:r><w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr><w:footnoteRef/></w:r>
                 # We insert this at the beginning of the paragraph element children
                 
                 ref_run = etree.Element(qn('w:r'))
                 ref_rPr = etree.SubElement(ref_run, qn('w:rPr'))
                 ref_style = etree.SubElement(ref_rPr, qn('w:rStyle'))
                 ref_style.set(qn('w:val'), 'FootnoteReference')
                 etree.SubElement(ref_run, qn('w:footnoteRef'))
                 
                 # Insert at correct position (after pPr if exists)
                 if len(wp) > 0 and wp[0].tag == qn('w:pPr'):
                     wp.insert(1, ref_run)
                 else:
                     wp.insert(0, ref_run)

                 updated_count += 1
                 
        if updated_count > 0:
            # Save back
            new_xml = etree.tostring(root, encoding='utf-8', standalone=True)
            footnotes_part._blob = new_xml
            print(f"Updated {updated_count} footnotes in footnotes.xml")
            
    except Exception as e:
        print(f"Error updating footnotes: {e}")
        import traceback
        traceback.print_exc()

def reassemble_docx(original_path: str, output_path: str, segments: List[SegmentInternal]):
    """
    Takes original DOCX and a list of TRANSLATED segments.
    Writes a new DOCX where text is replaced by target_content.
    """
    # 1. Copy file to output path
    shutil.copy(original_path, output_path)
    
    doc = Document(output_path)
    
    # We assume segments map 1:1 to paragraphs for this MVP.
    # In reality, we must match IDs or indices.
    # Our parser saved "original_index" in metadata!
    
    # segment_map caused collisions for split segments (make_key ignored sub_index).
    # We now pass the full list of segments to _inject_into_container and let it group them properly.
    
    # 1. Body
    _inject_into_container(doc, {"type": "body"}, segments)
    
    # 2. Sections
    for s_idx, section in enumerate(doc.sections):
        if section.header:
            _inject_into_container(section.header, {"type": "header", "section_index": s_idx}, segments)
        if section.footer:
             _inject_into_container(section.footer, {"type": "footer", "section_index": s_idx}, segments)

    # 3. Comments (Update comments.xml)
    _inject_comments(doc, segments)

    # 4. Footnotes
    _inject_footnotes(doc, segments)

    doc.save(output_path)

def _inject_into_container(container, base_metadata, source_segments):
    # source_segments: List[SegmentInternal]
    
    # Helper to merge segments targeting the same paragraph
    def get_merged_content(segs_for_para):
        if not segs_for_para:
            return "", {}
        full_text = ""
        combined_tags = {}
        # Sort by sub_index
        segs_for_para.sort(key=lambda x: x.metadata.get("sub_index", 0))
        for s in segs_for_para:
            text = s.target_content if s.target_content is not None else s.source_text
            # Restore whitespaces
            text = _restore_whitespaces(text, s.metadata)
            
            full_text += text
            if s.tags:
                combined_tags.update(s.tags)
                
        # Smart Merging Optimization
        # Remove </N><N> patterns (redundant boundary from split)
        # Regex: </(\d+)><\1>
        full_text = re.sub(r'</(\d+)><\1>', '', full_text)
        
        return full_text, combined_tags

    def _restore_whitespaces(text: str, metadata: dict) -> str:
        if not text or not metadata:
            return text
            
        ws = metadata.get("whitespaces")
        if not ws:
            return text
            
        leading = ws.get("leading", "")
        trailing = ws.get("trailing", "")
        
        # Restore Leading
        if leading:
             # Check if already present
             if not text.startswith(leading):
                 # If text starts with SOME whitespace, do we replace or append?
                 # Strategy: Ensure the EXACT whitespace exists.
                 # If text starts with stripped text, prepend.
                 # If text starts with different whitespace, maybe AI changed it?
                 # User wants to FIX truncation.
                 current_leading = re.match(r'^(\s+)', text)
                 if not current_leading:
                     text = leading + text
                 # If it has whitespace, but different? We trust AI or Source?
                 # User says "Model cuts them off".
                 # Providing the source leading space is usually safe for layout.
        
        # Restore Trailing
        if trailing:
             if not text.endswith(trailing):
                 current_trailing = re.search(r'(\s+)$', text)
                 if not current_trailing:
                     text = text + trailing
                     
        return text

    # 1. Build Grouped Segments from the flat list
    grouped_segments = {}
    
    for s in source_segments:
        m = s.metadata
        if not m: continue
        
        stype = m.get("type", "body")
        section_idx = m.get("section_index", -1)
        
        # Table Coords
        t_coords = None
        if m.get("child_type") == "table_cell" or m.get("sub_type") == "table" or stype == "table":
             t_coords = (m.get("table_index"), m.get("row_index"), m.get("cell_index"))
             
        p_idx = m.get("p_index") if m.get("p_index") is not None else m.get("index")
            
        key = (stype, section_idx, t_coords, p_idx)
        
        if key not in grouped_segments:
            grouped_segments[key] = []
        grouped_segments[key].append(s)

    # 2. Iterate the Container and Inject
    # We iterate the Physical Container (doc, header, etc.)
    
    stype = base_metadata.get("type")
    s_idx = base_metadata.get("section_index", -1)

    # DEBUG DUMP
    with open("debug_reassembly_keys.log", "a") as f:
         f.write(f"--- Injecting into {stype} s_idx={s_idx} ---\n")
         f.write(f"Grouped Keys: {[k for k in grouped_segments.keys() if k[0] == stype]}\n")

    # A. Paragraphs
    for i, para in enumerate(container.paragraphs):
        # Key = (stype, s_idx, None, i)
        key = (stype, s_idx, None, i)
        
        if key in grouped_segments:
            try:
                text, tags = get_merged_content(grouped_segments[key])
                # Restore whitespaces from the FIRST segment of the group (assuming 1:1 or first has info)
                # Actually, merged content might span multiple source segments?
                # Usually 1 paragraph = 1 or more segments.
                # If we merged them, `text` is the concatenation.
                # `grouped_segments[key]` is list of segments.
                # We should apply restoration to individual segments BEFORE merging?
                # No, get_merged_content merges TARGET content.
                # We need to apply whitespace logic PER SEGMENT.
                # But here we are injecting the whole paragraph.
                # Let's apply it using the metadata of the segments.
                
                # REVISION: `get_merged_content` concatenates raw target_contents.
                # If we modify `get_merged_content` to apply restoration internally loop?
                
                _inject_tagged_text(para, text, tags)
                with open("debug_reassembly_keys.log", "a") as f: f.write(f"Inject Para {key}: OK\n")
            except Exception as e:
                print(f"Error injecting segment {key}: {e}")
        else:
            with open("debug_reassembly_keys.log", "a") as f: f.write(f"Inject Para {key}: MISSED\n")

    # B. Tables
    for t_i, table in enumerate(container.tables):
        for r_i, row in enumerate(table.rows):
            for c_i, cell in enumerate(row.cells):
                for p_i, para in enumerate(cell.paragraphs):
                    # Key Construction
                    target_type = stype
                    if stype == "body":
                        target_type = "table"
                    
                    key = (target_type, s_idx, (t_i, r_i, c_i), p_i)
                    
                    if key in grouped_segments:
                        try:
                            text, tags = get_merged_content(grouped_segments[key])
                            _inject_tagged_text(para, text, tags)
                            with open("debug_reassembly_keys.log", "a") as f: f.write(f"Inject Table {key}: OK\n")
                        except Exception as e:
                            print(f"Error injecting table segment {key}: {e}")
                    else:
                        with open("debug_reassembly_keys.log", "a") as f: f.write(f"Inject Table {key}: MISSED\n")


    
# This requires major refactoring of the loops above.
# Let's do it in the next tool call properly.
def _inject_segment(para, segment):
    target_text = segment.target_content
    # Clear existing runs
    para.clear()
    # RE-INJECT
    _inject_tagged_text(para, target_text, segment.tags)

def _inject_tagged_text(paragraph, text, tags_map):
    """
    Parses 'text' which may contain:
    1. Custom Tags: <1>...</1> (mapped to properties via tags_map)
    2. HTML Tags: <b>, <strong>, <i>, <em>, <u>, <br/>
    3. Custom Markers: [TAB], [COMMENT], [SHAPE]
    
    Reconstructs the paragraph with appropriate runs and formatting.
    Preserves w:drawing and w:pict elements.
    """
    p_element = paragraph._element
    
    # 1. Preserve Shapes (Drawings/Picts) before clearing
    # We iterate descendants to find them in order.
    preserved_shapes = []
    # Note: element.iter() might return self? No, usually descendants.
    # We check specific tags.
    # Safest way to get in-order:
    for child in p_element.iter():
        if child.tag == qn('w:drawing') or child.tag == qn('w:pict'):
            try:
                # Deepcopy to detach and save
                preserved_shapes.append(copy.deepcopy(child))
            except Exception as e:
                print(f"Warning: Failed to preserve shape: {e}")

    # 2. Clear existing content
    p_element.clear_content()

    # DEBUG LOGGING
    with open("debug_reassembly.txt", "a") as f:
        f.write(f"\n--- Injecting Text ---\nText: {text[:50]}...\nTags Map Keys: {list(tags_map.keys())}\n")
        f.write(f"Preserved Shapes: {len(preserved_shapes)}\n")

    # Tokenize
    tokens = re.split(r'(<[^>]+>)', text)
    
    active_style = {'bold': 0, 'italic': 0, 'underline': 0, 'highlight': False, 'superscript': False, 'subscript': False}
    
    for token in tokens:
        if not token:
            continue
            
        # Is it a Tag?
        if token.startswith("<") and token.endswith(">"):
            tag_content = token[1:-1] # strip < >
            is_closing = tag_content.startswith("/")
            if is_closing:
                tag_content = tag_content[1:]
                
            # Check ID (Digits) -> Custom Tag
            if tag_content.isdigit():
                tid = tag_content
                tag = tags_map.get(tid)
                if tag:
                    # Apply tag formatting
                    if tag.type == 'bold':
                        active_style['bold'] += -1 if is_closing else 1
                    elif tag.type == 'italic':
                        active_style['italic'] += -1 if is_closing else 1
                    elif tag.type == 'underline':
                        active_style['underline'] += -1 if is_closing else 1
                    elif tag.type == 'superscript':
                        active_style['superscript'] = not is_closing
                    elif tag.type == 'subscript':
                        active_style['subscript'] = not is_closing
                    elif tag.type == 'color':
                        # Handle color application
                        if is_closing:
                             active_style.pop('color', None)
                        else:
                             if tag.xml_attributes and 'color' in tag.xml_attributes:
                                 active_style['color'] = tag.xml_attributes['color']
                                 
                    elif tag.type == 'comment':
                        active_style['highlight'] = not is_closing
                    elif tag.type == 'highlight':
                        if is_closing:
                             active_style['highlight'] = False
                        else:
                             # Store specific color if available
                             if tag.xml_attributes and 'color' in tag.xml_attributes:
                                 active_style['highlight'] = tag.xml_attributes['color']
                             else:
                                 active_style['highlight'] = True
                    elif tag.type == 'shape' and not is_closing:
                        # Insert Shape
                        if preserved_shapes:
                            shape_el = preserved_shapes.pop(0)
                            # Shape must be in a run
                            run = paragraph.add_run()
                            run._element.append(shape_el)
                        else:
                            # Shape missing/deleted?
                            # Create a placeholder or ignore
                            run = paragraph.add_run("[MISSING SHAPE]")
                            run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)
                            
                    # Reconstruct Hyperlinks
                    elif tag.type == 'link':
                         if is_closing:
                             # Close hyperlink
                             if 'hyperlink_el' in active_style:
                                 active_style.pop('hyperlink_el')
                         else:
                             # Open Hyperlink
                             # 1. Create w:hyperlink
                             hyplink = docx.oxml.shared.OxmlElement('w:hyperlink')
                             if tag.xml_attributes and 'rid' in tag.xml_attributes:
                                 hyplink.set(qn('r:id'), tag.xml_attributes['rid'])
                             
                             # 2. Append to paragraph
                             paragraph._element.append(hyplink)
                             
                             # 3. Set context
                             active_style['hyperlink_el'] = hyplink
                        
            # Check HTML Tags
            else:
                lower_tag = tag_content.lower()
                if lower_tag == 'br/':
                    paragraph.add_run().add_break()
                    continue
                elif lower_tag in ['b', 'strong']:
                   active_style['bold'] += -1 if is_closing else 1
                elif lower_tag in ['i', 'em']:
                   active_style['italic'] += -1 if is_closing else 1
                elif lower_tag == 'u':
                   active_style['underline'] += -1 if is_closing else 1
                     
        else:
            # Generic Text
            
            # Helper to add run
            def add_styled_run(content):
                # context: where to add run?
                # If inside hyperlink, add to hyperlink_el
                if 'hyperlink_el' in active_style:
                    parent = active_style['hyperlink_el']
                    run = docx.oxml.shared.OxmlElement('w:r')
                    t = docx.oxml.shared.OxmlElement('w:t')
                    t.text = content
                    # Space preserve
                    if len(content.strip()) < len(content):
                         t.set(qn('xml:space'), 'preserve')
                    run.append(t)
                    parent.append(run)
                    
                    # Apply styles to this new run
                    # We need to construct rPr
                    rPr = docx.oxml.shared.OxmlElement('w:rPr')
                    has_style = False
                    
                    if active_style['bold'] > 0: 
                        rPr.append(docx.oxml.shared.OxmlElement('w:b'))
                        has_style = True
                    if active_style['italic'] > 0: 
                        rPr.append(docx.oxml.shared.OxmlElement('w:i'))
                        has_style = True
                    if active_style['underline'] > 0: 
                        u = docx.oxml.shared.OxmlElement('w:u')
                        u.set(qn('w:val'), 'single')
                        rPr.append(u)
                        has_style = True
                    
                    # Hyperlink style (usually Blue + Underline)
                    # Use existing style if defined? "Hyperlink"
                    rStyle = docx.oxml.shared.OxmlElement('w:rStyle')
                    rStyle.set(qn('w:val'), 'Hyperlink')
                    rPr.append(rStyle)
                    has_style = True
                        
                    if has_style:
                        run.append(rPr)

                else:
                    # Normal Run
                    run = paragraph.add_run(content)
                    if active_style['bold'] > 0: run.bold = True
                    if active_style['italic'] > 0: run.italic = True
                    if active_style['underline'] > 0: run.underline = True
                    if active_style['superscript']: run.font.superscript = True
                    if active_style['subscript']: run.font.subscript = True
                    
                    if 'color' in active_style:
                         try:
                             run.font.color.rgb = docx.shared.RGBColor.from_string(active_style['color'])
                         except:
                             pass 
                             
                    if active_style['highlight']: 
                         # Parse highlight color or default to yellow
                         hl_color = active_style['highlight']
                         if isinstance(hl_color, str) and hl_color != 'True' and hl_color != '1' and hl_color != 'True':
                              try:
                                   if hasattr(WD_COLOR_INDEX, hl_color):
                                        run.font.highlight_color = getattr(WD_COLOR_INDEX, hl_color)
                                   else:
                                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                              except:
                                   run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                         else:
                              run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            
            # Handle [TAB]
            if "[TAB]" in token:
                parts = token.split("[TAB]")
                for i, part in enumerate(parts):
                    if i > 0:
                        paragraph.add_run().add_tab()
                    if part:
                         add_styled_run(part)
            else:
                add_styled_run(token)

    # Append remaining shapes if any (user deleted tag?)
    if preserved_shapes:
        for shape_el in preserved_shapes:
             run = paragraph.add_run()
             run._element.append(shape_el)


