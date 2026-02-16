import re
import copy
import docx
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.text.paragraph import Paragraph
from app.logger import get_logger

logger = get_logger("Assembler")

MC_NS = 'http://schemas.openxmlformats.org/markup-compatibility/2006'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def _parse_tag_token(token: str):
    """Parse a tag token like '<3>' or '</3>' into (tag_content, is_closing)."""
    tag_content = token[1:-1]
    is_closing = tag_content.startswith("/")
    if is_closing:
        tag_content = tag_content[1:]
    return tag_content, is_closing


def _parse_tc_attrs(tag_content: str) -> dict:
    """Extract data-op-* attributes from a TC tag like 'insert data-op-user-id="x" ...'."""
    attrs = {}
    for m in re.finditer(r'data-op-([\w-]+)="([^"]*)"', tag_content):
        attrs[m.group(1)] = m.group(2)
    return attrs


def _format_tc_date(date_str: str) -> str:
    """Convert a date string to DOCX-compatible ISO 8601 format.

    Handles:
    - ISO 8601: "2024-01-15T10:00:00Z" → pass through
    - Space-separated: "2024-01-15 10:00" → "2024-01-15T10:00:00Z"
    - Unix timestamp ms (from tiptap getMinuteTime()): "1708099200000" → ISO
    - Empty/missing → current UTC time as fallback
    """
    from datetime import datetime as _dt

    if not date_str:
        return _dt.utcnow().strftime("%Y-%m-%dT%H:%M:00Z")
    if "T" in date_str and date_str.endswith("Z"):
        return date_str
    try:
        date_str = date_str.strip()
        # Unix timestamp in milliseconds (from tiptap track-change-extension)
        if date_str.isdigit() and len(date_str) >= 10:
            ts = int(date_str) / 1000
            return _dt.utcfromtimestamp(ts).strftime("%Y-%m-%dT%H:%M:00Z")
        if " " in date_str and "T" not in date_str:
            date_str = date_str.replace(" ", "T", 1)
        if not date_str.endswith("Z"):
            parts = date_str.split("T")
            if len(parts) == 2:
                time_part = parts[1]
                if time_part.count(":") < 2:
                    time_part += ":00"
                date_str = parts[0] + "T" + time_part + "Z"
            else:
                date_str += "T00:00:00Z"
        return date_str
    except Exception:
        return _dt.utcnow().strftime("%Y-%m-%dT%H:%M:00Z")


class AssemblerContext:
    """
    Encapsulates the state and handlers for injecting tagged text into a paragraph.

    Tag handlers are registered in TAG_HANDLERS. To add a new tag type:
    1. Add an entry mapping type name -> handler method name
    2. Implement the handler method
    """

    TAG_HANDLERS = {
        'bold': '_handle_toggle',
        'italic': '_handle_toggle',
        'underline': '_handle_toggle',
        'superscript': '_handle_flag',
        'subscript': '_handle_flag',
        'strike': '_handle_flag',
        'smallCaps': '_handle_flag',
        'caps': '_handle_flag',
        'color': '_handle_valued_prop',
        'size': '_handle_valued_prop',
        'font': '_handle_valued_prop',
        'comment': '_handle_highlight',
        'highlight': '_handle_highlight',
        'link': '_handle_link',
        'shape': '_handle_shape',
        'footnote': '_handle_note_ref',
        'endnote': '_handle_note_ref',
    }

    # Class-level counter for unique TC revision IDs across paragraphs
    _tc_rev_counter = 10000

    def __init__(self, paragraph: Paragraph, tags_map: dict, shape_map=None):
        self.paragraph = paragraph
        self.p_element = paragraph._element
        self.tags_map = tags_map
        self.shape_map = shape_map
        self.preserved_shapes = []
        self.active_style = {
            'bold': 0, 'italic': 0, 'underline': 0,
            'highlight': False,
            'superscript': False, 'subscript': False,
            'strike': False, 'smallCaps': False, 'caps': False,
        }
        # Track Changes state: w:ins / w:del wrapper element
        self.tc_wrapper = None   # Current OxmlElement (w:ins or w:del)
        self.tc_type = None      # 'insert' or 'delete'

    def run(self, text: str):
        """Main entry point: preserve shapes, clear, inject tokens, append remaining shapes."""
        self.preserve_shapes()
        self.clear_paragraph()

        tokens = re.split(r'(<[^>]+>)', text)
        for token in tokens:
            if not token:
                continue
            if token.startswith("<") and token.endswith(">"):
                self._handle_tag_token(token)
            else:
                self.add_styled_run(token)

        self.append_remaining_shapes()

    # --- Setup / Teardown ---

    def preserve_shapes(self):
        """Scan paragraph for drawing/pict elements and deep-copy them before clearing."""
        for child in self.p_element.iter():
            if child.tag == f"{{{MC_NS}}}AlternateContent":
                try:
                    self.preserved_shapes.append(copy.deepcopy(child))
                except Exception as e:
                    logger.warning(f"Failed to preserve AlternateContent: {e}")

            elif child.tag == qn('w:drawing') or child.tag == qn('w:pict'):
                parent = child.getparent()
                is_wrapped = False
                if parent is not None:
                    if parent.tag == f"{{{MC_NS}}}Choice" or parent.tag == f"{{{MC_NS}}}Fallback":
                        is_wrapped = True
                if not is_wrapped:
                    try:
                        self.preserved_shapes.append(copy.deepcopy(child))
                    except Exception as e:
                        logger.warning(f"Failed to preserve shape: {e}")

    def clear_paragraph(self):
        """Remove all children except w:pPr (paragraph properties)."""
        for child in list(self.p_element):
            if child.tag == qn('w:pPr'):
                continue
            self.p_element.remove(child)

    def append_remaining_shapes(self):
        """Append any shapes not consumed by explicit <shape> tags (e.g. pure images)."""
        for remaining_shape in self.preserved_shapes:
            run = self.paragraph.add_run()
            run._element.append(remaining_shape)

    # --- Token Dispatch ---

    def _handle_tag_token(self, token: str):
        tag_content, is_closing = _parse_tag_token(token)

        if tag_content.isdigit():
            tid = tag_content
            tag = self.tags_map.get(tid)
            if not tag:
                return
            handler_name = self.TAG_HANDLERS.get(tag.type)
            if handler_name:
                getattr(self, handler_name)(tid, tag, is_closing)
        else:
            self._handle_html_tag(tag_content, is_closing)

    # --- Tag Handlers ---

    def _handle_toggle(self, tid, tag, is_closing):
        """Handle counter-based toggles: bold, italic, underline."""
        self.active_style[tag.type] += -1 if is_closing else 1

    def _handle_flag(self, tid, tag, is_closing):
        """Handle boolean flags: superscript, subscript, strike, smallCaps, caps."""
        self.active_style[tag.type] = not is_closing

    def _handle_valued_prop(self, tid, tag, is_closing):
        """Handle valued properties: color, size, font."""
        if is_closing:
            self.active_style.pop(tag.type, None)
        else:
            attrs = tag.xml_attributes or {}
            if tag.type == 'color':
                val = attrs.get('color')
            elif tag.type == 'size':
                val = attrs.get('val')
            elif tag.type == 'font':
                val = attrs.get('name')
            else:
                val = None
            if val:
                self.active_style[tag.type] = val

    def _handle_highlight(self, tid, tag, is_closing):
        """Handle comment and highlight tags (both map to highlight style)."""
        if tag.type == 'comment':
            self.active_style['highlight'] = not is_closing
        else:
            self.active_style['highlight'] = False if is_closing else True

    def _handle_link(self, tid, tag, is_closing):
        """Handle hyperlink open/close."""
        if is_closing:
            self.active_style.pop('hyperlink_el', None)
        else:
            hyplink = docx.oxml.shared.OxmlElement('w:hyperlink')
            if tag.xml_attributes and 'rid' in tag.xml_attributes:
                hyplink.set(qn('r:id'), tag.xml_attributes['rid'])
            self.paragraph._element.append(hyplink)
            self.active_style['hyperlink_el'] = hyplink

    def _handle_shape(self, tid, tag, is_closing):
        """Handle shape tags (only opening — shapes are self-contained)."""
        if is_closing:
            return
        if self.preserved_shapes:
            shape_el = self.preserved_shapes.pop(0)
            if self.shape_map and tag.xml_attributes and 'id' in tag.xml_attributes:
                sid = tag.xml_attributes['id']
                self._process_shape_element(shape_el, sid)
            run = self.paragraph.add_run()
            run._element.append(shape_el)
        else:
            run = self.paragraph.add_run("[MISSING SHAPE]")
            run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)

    def _handle_note_ref(self, tid, tag, is_closing):
        """Handle footnote/endnote reference tags (only opening)."""
        if is_closing:
            return
        run = self.paragraph.add_run()
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
        rStyle = docx.oxml.shared.OxmlElement('w:rStyle')

        if tag.type == 'footnote':
            rStyle.set(qn('w:val'), 'FootnoteReference')
            rPr.append(rStyle)
            run._element.append(rPr)
            ref = docx.oxml.shared.OxmlElement('w:footnoteReference')
            ref.set(qn('w:id'), tag.xml_attributes['id'])
            run._element.append(ref)
        else:  # endnote
            rStyle.set(qn('w:val'), 'EndnoteReference')
            rPr.append(rStyle)
            run._element.append(rPr)
            ref = docx.oxml.shared.OxmlElement('w:endnoteReference')
            ref.set(qn('w:id'), tag.xml_attributes['id'])
            run._element.append(ref)

    def _handle_html_tag(self, tag_content: str, is_closing: bool):
        """Handle residual HTML tags (br, b, i, u) and TC tags (insert, delete)."""
        # Extract tag name (first word) for matching
        tag_name = tag_content.split(None, 1)[0].lower() if tag_content else ""

        if tag_name == 'br/':
            self.paragraph.add_run().add_break()
        elif tag_name in ['b', 'strong']:
            self.active_style['bold'] += -1 if is_closing else 1
        elif tag_name in ['i', 'em']:
            self.active_style['italic'] += -1 if is_closing else 1
        elif tag_name == 'u':
            self.active_style['underline'] += -1 if is_closing else 1
        elif tag_name == 'insert':
            self._handle_tc_tag(tag_content, is_closing, 'insert')
        elif tag_name == 'delete':
            self._handle_tc_tag(tag_content, is_closing, 'delete')

    def _handle_tc_tag(self, tag_content: str, is_closing: bool, tc_type: str):
        """Handle Track Changes <insert>/<delete> tags → w:ins/w:del XML."""
        if is_closing:
            self.tc_wrapper = None
            self.tc_type = None
            return

        attrs = _parse_tc_attrs(tag_content)
        author = attrs.get('user-nickname', attrs.get('user-id', 'Unknown'))
        date = _format_tc_date(attrs.get('date', ''))

        AssemblerContext._tc_rev_counter += 1
        rev_id = str(AssemblerContext._tc_rev_counter)

        w_tag = 'w:ins' if tc_type == 'insert' else 'w:del'
        tc_el = OxmlElement(w_tag)
        tc_el.set(qn('w:id'), rev_id)
        tc_el.set(qn('w:author'), author)
        tc_el.set(qn('w:date'), date)
        self.p_element.append(tc_el)

        self.tc_wrapper = tc_el
        self.tc_type = tc_type

    # --- Run Creation ---

    def add_styled_run(self, content: str):
        """Create a run with the current active formatting style."""
        if self.tc_wrapper is not None:
            self._add_tc_run(content)
        elif 'hyperlink_el' in self.active_style:
            self._add_hyperlink_run(content)
        else:
            self._add_paragraph_run(content)

    def _add_tc_run(self, content: str):
        """Add a run inside a w:ins or w:del element."""
        run_el = OxmlElement('w:r')

        # Run properties (formatting)
        rPr = self._build_rpr()
        if rPr is not None:
            run_el.append(rPr)

        # w:del uses w:delText, w:ins uses w:t
        if self.tc_type == 'delete':
            t = OxmlElement('w:delText')
        else:
            t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = content
        run_el.append(t)

        self.tc_wrapper.append(run_el)

    def _build_rpr(self):
        """Build a w:rPr element from the current active_style. Returns None if no formatting."""
        parts = []
        if self.active_style['bold'] > 0:
            parts.append(OxmlElement('w:b'))
        if self.active_style['italic'] > 0:
            parts.append(OxmlElement('w:i'))
        if self.active_style['underline'] > 0:
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            parts.append(u)
        if self.active_style.get('superscript'):
            vAlign = OxmlElement('w:vertAlign')
            vAlign.set(qn('w:val'), 'superscript')
            parts.append(vAlign)
        if self.active_style.get('subscript'):
            vAlign = OxmlElement('w:vertAlign')
            vAlign.set(qn('w:val'), 'subscript')
            parts.append(vAlign)
        if self.active_style.get('strike'):
            parts.append(OxmlElement('w:strike'))
        if self.active_style.get('smallCaps'):
            parts.append(OxmlElement('w:smallCaps'))
        if self.active_style.get('caps'):
            parts.append(OxmlElement('w:caps'))
        if 'color' in self.active_style:
            c = OxmlElement('w:color')
            c.set(qn('w:val'), self.active_style['color'])
            parts.append(c)
        if 'size' in self.active_style:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), self.active_style['size'])
            parts.append(sz)
        if 'font' in self.active_style:
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), self.active_style['font'])
            rFonts.set(qn('w:hAnsi'), self.active_style['font'])
            parts.append(rFonts)
        if not parts:
            return None
        rPr = OxmlElement('w:rPr')
        for p in parts:
            rPr.append(p)
        return rPr

    def _add_hyperlink_run(self, content: str):
        """Add a run inside the current active hyperlink element."""
        parent = self.active_style['hyperlink_el']

        run = docx.oxml.shared.OxmlElement('w:r')
        t = docx.oxml.shared.OxmlElement('w:t')
        if content:
            t.set(qn('xml:space'), 'preserve')
        t.text = content

        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        # Hyperlink Style
        rStyle = docx.oxml.shared.OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'Hyperlink')
        rPr.append(rStyle)

        color = docx.oxml.shared.OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')
        rPr.append(color)

        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

        # Apply other active styles
        if self.active_style['bold'] > 0:
            rPr.append(docx.oxml.shared.OxmlElement('w:b'))
        if self.active_style['italic'] > 0:
            rPr.append(docx.oxml.shared.OxmlElement('w:i'))

        run.append(rPr)
        run.append(t)
        parent.append(run)

    def _add_paragraph_run(self, content: str):
        """Add a standard styled run to the paragraph."""
        run = self.paragraph.add_run(content)

        if self.active_style['bold'] > 0: run.bold = True
        if self.active_style['italic'] > 0: run.italic = True
        if self.active_style['underline'] > 0: run.underline = True

        if self.active_style['superscript']:
            run.font.superscript = True
        if self.active_style['subscript']:
            run.font.subscript = True
        if self.active_style['strike']:
            run.font.strike = True
        if self.active_style['smallCaps']:
            run.font.small_caps = True
        if self.active_style['caps']:
            run.font.all_caps = True

        if 'color' in self.active_style:
            try:
                run.font.color.rgb = docx.shared.RGBColor.from_string(self.active_style['color'])
            except: pass

        if 'size' in self.active_style:
            try:
                pt_val = int(self.active_style['size']) / 2
                run.font.size = docx.shared.Pt(pt_val)
            except: pass

        if 'font' in self.active_style:
            run.font.name = self.active_style['font']

        if self.active_style['highlight'] is not False:
            if self.active_style['highlight'] is True:
                run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
            elif isinstance(self.active_style['highlight'], str):
                run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW

    # --- Shape Processing ---

    def _process_shape_element(self, element, sid):
        """Recursively inject translated text into textbox content of a shape element."""
        ns = {'w': W_NS}
        txbx_contents = element.findall('.//w:txbxContent', ns)

        global_p_idx = 0
        for txbx in txbx_contents:
            paras = txbx.findall('.//w:p', ns)
            for para in paras:
                if sid in self.shape_map:
                    t_segs = self.shape_map[sid]
                    matching_segs = [s for s in t_segs if s.metadata.get("p_index") == global_p_idx]

                    if matching_segs:
                        matching_segs.sort(key=lambda x: x.metadata.get("sub_index", 0))

                        full_text = ""
                        combined_tags = {}
                        for s in matching_segs:
                            content = s.target_content if s.target_content is not None else s.source_text
                            full_text += content
                            if s.tags:
                                combined_tags.update(s.tags)

                        full_text = re.sub(r'</(\d+)><\1>', '', full_text)

                        proxy_p = Paragraph(para, None)
                        inject_tagged_text(proxy_p, full_text, combined_tags, None)

                global_p_idx += 1

        # Handle AlternateContent wrapper
        if element.tag == f"{{{MC_NS}}}AlternateContent":
            mc_ns = {'mc': MC_NS}
            choice = element.find('mc:Choice', mc_ns)
            if choice is not None:
                for child in choice:
                    self._process_shape_element(child, sid)
            fallback = element.find('mc:Fallback', mc_ns)
            if fallback is not None:
                for child in fallback:
                    self._process_shape_element(child, sid)


def inject_tagged_text(paragraph: Paragraph, text: str, tags_map: dict, shape_map=None):
    """
    Parses 'text' containing custom tags and reconstructs the paragraph with formatting.
    Preserves w:drawing and w:pict elements found in the paragraph before clearing.
    """
    ctx = AssemblerContext(paragraph, tags_map, shape_map)
    ctx.run(text)
