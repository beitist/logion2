from typing import List, Tuple, Dict
import uuid
import re
import docx
from docx.oxml.ns import qn
from docx.api import Document
from docx.text.run import Run
from lxml import etree

from .schemas import SegmentInternal, TagModel
from .logger import get_logger

logger = get_logger("Parser")

def parse_docx(file_path: str, segmentation_func=None, source_lang="en") -> List[SegmentInternal]:
    """
    Parses a DOCX file and extracts segments with tags for formatting, hyperlinks, and comments.
    """
    doc = Document(file_path)
    segments = []
    
    # 1. Load Comments Map (id -> text)
    comments_map = {}
    try:
        # Try explicit XML parsing for comments as python-docx support varies
        # We need to find the comments part
        part = doc.part
        # This is a bit hacky but robust for reading: iterate relations
        for rel in part.rels.values():
            if "comments" in rel.reltype and not "commentsExtended" in rel.reltype and not "commentsIds" in rel.reltype:
                 # Found comments.xml
                 xml_data = rel.target_part.blob
                 root = etree.fromstring(xml_data)
                 # Namespace usually w:
                 namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                 for comment in root.findall('.//w:comment', namespaces):
                     cid = comment.get(qn('w:id'))
                     ctext = "".join([t.text for t in comment.findall('.//w:t', namespaces) if t.text])
                     comments_map[cid] = ctext
    except Exception as e:
        logger.warning(f"Could not load comments: {e}")

    # Helper context to pass comments_map, segmentation_func, source_lang
    context = {
        "comments_map": comments_map,
        "segmentation_func": segmentation_func,
        "segmentation_func": segmentation_func,
        "source_lang": source_lang,
        "extra_segments": []
    }

    # 1. Body
    segments.extend(_process_container(doc, {"type": "body"}, context))

    # 2. Key Sections (Headers/Footers)
    for s_idx, section in enumerate(doc.sections):
        # Header
        if section.header:
             segments.extend(_process_container(section.header, {
                 "type": "header", 
                 "section_index": s_idx
             }, context))
        # Footer
        if section.footer:
             segments.extend(_process_container(section.footer, {
                 "type": "footer", 
                 "section_index": s_idx
             }, context))
             
    # 3. Comments (Separate Segments)
    # Strategy: Expose comments as translatable segments.
    # We use valid IDs from the map.
    # Note: This list might include comments not referenced in the document body if they are zombies,
    # but that's acceptable.
    for cid, ctext in comments_map.items():
        if not ctext.strip():
            continue
            
        # Create a SegmentInternal for the comment
        # Metadata distinguishes it
        seg = SegmentInternal(
            id=str(uuid.uuid4()),
            segment_id=str(uuid.uuid4()),
            source_text=ctext,
            target_content=None,
            status="draft",
            tags={},
            metadata={
                "type": "comment",
                "comment_id": cid
            }
        )
        segments.append(seg)

    # 4. Footnotes
    segments.extend(_extract_footnotes(doc, context))
    
    # 5. Endnotes
    segments.extend(_extract_endnotes(doc, context))
    
    # 6. Extra Segments (Shapes/Textboxes)
    if "extra_segments" in context:
         segments.extend(context["extra_segments"])

    return segments

def _extract_footnotes(doc, context) -> List[SegmentInternal]:
    """
    Extracts footnotes as separate segments.
    """
    segments = []
    try:
        part = doc.part
        footnotes_part = None
        for rel in part.rels.values():
             if "footnotes" in rel.reltype:
                 footnotes_part = rel.target_part
                 break
                 
        if not footnotes_part:
            return []
            
        xml_data = footnotes_part.blob
        root = etree.fromstring(xml_data)
        namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        for footnote in root.findall('.//w:footnote', namespaces):
            fid = footnote.get(qn('w:id'))
            ftype = footnote.get(qn('w:type'))
            if ftype in ["separator", "continuationSeparator"]:
                continue

            for i, p_elem in enumerate(footnote.findall('.//w:p', namespaces)):
                # Manual XML processing -> _process_paragraph expects XML now?
                # or we convert p_elem to XML wrapper?
                # Let's update _process_paragraph to take XML element directly.
                
                # Check text content via xpath
                text_content = "".join([t.text for t in p_elem.findall('.//w:t', namespaces) if t.text])
                if not text_content.strip():
                     continue

                meta = {
                    "type": "footnote",
                    "footnote_id": fid,
                    "p_index": i
                }
                
                found_segments = _process_paragraph(p_elem, meta, context)
                segments.extend(found_segments)
            
    except Exception as e:
        logger.warning(f"Could not load footnotes: {e}")
        
    return segments

def _extract_endnotes(doc, context) -> List[SegmentInternal]:
    """
    Extracts endnotes as separate segments.
    """
    segments = []
    try:
        part = doc.part
        endnotes_part = None
        for rel in part.rels.values():
             if "endnotes" in rel.reltype:
                 endnotes_part = rel.target_part
                 break
                 
        if not endnotes_part:
            return []
            
        xml_data = endnotes_part.blob
        root = etree.fromstring(xml_data)
        namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        for endnote in root.findall('.//w:endnote', namespaces):
            eid = endnote.get(qn('w:id'))
            etype = endnote.get(qn('w:type'))
            if etype in ["separator", "continuationSeparator"]:
                continue

            for i, p_elem in enumerate(endnote.findall('.//w:p', namespaces)):
                
                text_content = "".join([t.text for t in p_elem.findall('.//w:t', namespaces) if t.text])
                if not text_content.strip():
                     continue

                meta = {
                    "type": "endnote",
                    "endnote_id": eid,
                    "p_index": i
                }
                
                segments.extend(_process_paragraph(p_elem, meta, context))
            
    except Exception as e:
        logger.warning(f"Could not load endnotes: {e}")
        
    return segments

def _process_container(container, base_metadata: dict, context: dict) -> List[SegmentInternal]:
    container_segments = []
    
    # 1. Paragraphs
    for i, para in enumerate(container.paragraphs):
        # REMOVED: if not para.text.strip(): continue (Skips shapes/images!)
        
        # Merge base_metadata with specific location
        loc = base_metadata.copy()
        # "index" was used for body paragraphs in WP1. 
        if base_metadata.get("type") == "body":
            loc["index"] = i
        else:
            loc["p_index"] = i
            
        segment_list = _process_paragraph(para._element, loc, context)
        if segment_list:
            container_segments.extend(segment_list)

    # 2. Tables
    for t_idx, table in enumerate(container.tables):
        # Track processed cells to handle Merged Cells (vMerge)
        # python-docx repeats the cell object for each covered row.
        processed_tcs = set() 
        
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                # Check for duplication via XML element (tc)
                if cell._tc in processed_tcs:
                    continue
                processed_tcs.add(cell._tc)

                # Recursion! A cell is also a container (has paragraphs and tables)
                # We need to construct the location correctly.
                
                cell_loc = base_metadata.copy()
                cell_loc["child_type"] = "table_cell" # differentiating direct paragraphs vs table cell paragraphs
                cell_loc["table_index"] = t_idx
                cell_loc["row_index"] = r_idx
                cell_loc["cell_index"] = c_idx
                
                # Cells only contain paragraphs/tables, so we can use _process_container?
                # But _process_container iterates tables too, which supports nested tables automatically!
                # However, our metadata scheme in WP1 Extension 1 was flat for tables: "type": "table".
                # To support recursion and headers properly, we need a flexible metadata structure.
                # For this step, let's keep it compatible but support the hierarchy.
                
                # Current table logic used:
                # location = { "type": "table", "table_index": ..., ... "p_index": ... }
                # The "type" was "table".
                
                # If we are in the body, base_metadata is {"type": "body"}.
                # If we are in header, base_metadata is {"type": "header", "section": ...}.
                
                # Let's adjust the location logic for specific paragraphs inside the cell.
                # We will NOT call _process_container recursively yet to avoid breaking previous schema too much,
                # but we will iterate the cell's paragraphs.
                
                for p_idx, para in enumerate(cell.paragraphs):
                    # REMOVED: if not para.text.strip(): continue
                        
                    loc = base_metadata.copy()
                    # We override "type" to indicate it's in a table, OR we keep parent type?
                    # The previous verification expected "type": "table".
                    # Let's keep "type": "table" but add "parent_context" if needed?
                    # Actually, for Headers, we need to know it is in a header.
                    
                    if base_metadata["type"] == "body":
                         loc["type"] = "table" # Backwards compatibility with previous step
                    else:
                         # It's a table inside a header/footer
                         loc["sub_type"] = "table"
                         
                    loc["table_index"] = t_idx
                    loc["row_index"] = r_idx
                    loc["cell_index"] = c_idx
                    loc["p_index"] = p_idx
                    
                    segment_list = _process_paragraph(para._element, loc, context)
                    if segment_list:
                        container_segments.extend(segment_list)

    return container_segments

def _process_paragraph(para_element, location: dict, context: dict) -> List[SegmentInternal]:
    """
    Converts a docx Paragraph XML ELEMENT into a SegmentInternal with tags.
    Handles Runs, Hyperlinks, and Comments via XML iteration.
    Argument 'para_element' is an lxml/OxmlElement (not Paragraph object).
    """
    full_text = ""
    tags = {}
    
    # Counter for tags
    tag_counter = 1
    
    def add_tag(tag_model: TagModel):
        nonlocal tag_counter
        tid = str(tag_counter)
        tags[tid] = tag_model
        tag_counter += 1
        return tid

    # IMPORTANT: We iterate DIRECT children of the paragraph element.
    # Hyperlinks and Runs are children of the paragraph.
    # We must treat them sequentially.
    
    run_buffer = []

    def flush_run_buffer():
        nonlocal run_buffer, full_text
        if not run_buffer: return
        
        # Merge Logic by Signature
        groups = []
        if run_buffer:
            current_group = [run_buffer[0]]
            current_sig = _get_run_signature(run_buffer[0])
            
            for r in run_buffer[1:]:
                sig = _get_run_signature(r)
                if sig == current_sig:
                    current_group.append(r)
                else:
                    groups.append(current_group)
                    current_group = [r]
                    current_sig = sig
            groups.append(current_group)
            
            for group in groups:
                 # Combined text
                merged_text = ""
                for r in group:
                    merged_text += _get_run_text(r, namespaces)
                
                if merged_text:
                    # Use properties from the FIRST element of the group
                    # But pass the merged text content
                    res = _process_run_element(group[0], add_tag, context, text_override=merged_text)
                    full_text += res

        run_buffer = []

    namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    for child in para_element:
        tag_name = child.tag
        
        # 1. Regular Run (w:r) - Check for Pure Text
        if tag_name == qn('w:r'):
             if _is_pure_text_run(child):
                 run_buffer.append(child)
                 continue
             else:
                 flush_run_buffer()
                 # Handle complex run immediately
                 run_text = _process_run_element(child, add_tag, context)
                 if run_text:
                     full_text += run_text
                 continue

        # For any other tag, flush buffer first
        flush_run_buffer()

        # 2. Hyperlink (w:hyperlink)
        if tag_name == qn('w:hyperlink'):
            rid = child.get(qn('r:id'))
            
            link_tag = TagModel(type="link", xml_attributes={"is_hyperlink": True, "rid": rid})
            tid = add_tag(link_tag)
            
            full_text += f"<{tid}>"
            
            # Iterate children of hyperlink (runs)
            for sub_child in child:
                if sub_child.tag == qn('w:r'):
                    # Treat as run
                    run_text = _process_run_element(sub_child, add_tag, context)
                    if run_text:
                        full_text += run_text
            
            full_text += f"</{tid}>"

        # 2b. Comment Range Start
        elif tag_name == qn('w:commentRangeStart'):
            comment_id = child.get(qn('w:id'))
            if comment_id and context["comments_map"].get(comment_id):
                comment_text = context["comments_map"][comment_id]
                com_tag = TagModel(type="comment", content=comment_text, ref_id=comment_id)
                tid = add_tag(com_tag)
                if "_active_ranges" not in context:
                    context["_active_ranges"] = {}
                context["_active_ranges"][comment_id] = tid
                full_text += f"<{tid}>"

        # 2c. Comment Range End
        elif tag_name == qn('w:commentRangeEnd'):
            comment_id = child.get(qn('w:id'))
            if "_active_ranges" in context and comment_id in context["_active_ranges"]:
                tid = context["_active_ranges"][comment_id]
                full_text += f"</{tid}>"
                if "_handled_ranges" not in context:
                    context["_handled_ranges"] = set()
                context["_handled_ranges"].add(comment_id)
                del context["_active_ranges"][comment_id]

        # 3. Comment Reference (w:commentReference)
        elif tag_name == qn('w:commentReference'):
            comment_id = child.get(qn('w:id'))
            was_handled = "_handled_ranges" in context and comment_id in context["_handled_ranges"]
            is_active = "_active_ranges" in context and comment_id in context["_active_ranges"]
            if was_handled or is_active:
                pass
            elif comment_id and context["comments_map"].get(comment_id):
                comment_text = context["comments_map"][comment_id]
                com_tag = TagModel(type="comment", content=comment_text, ref_id=comment_id)
                tid = add_tag(com_tag)
                full_text += f"<{tid}></{tid}>"

        # 4. Inserted Text (w:ins)
        elif tag_name == qn('w:ins'):
            for sub_child in child:
                if sub_child.tag == qn('w:r'):
                     run_text = _process_run_element(sub_child, add_tag, context)
                     if run_text:
                         full_text += run_text

        # 5. Footnote Reference
        elif tag_name == qn('w:footnoteReference'):
            fid = child.get(qn('w:id'))
            ftag = TagModel(type="footnote", xml_attributes={"id": fid})
            tid = add_tag(ftag)
            full_text += f"<{tid}></{tid}>"

        # 6. Endnote Reference
        elif tag_name == qn('w:endnoteReference'):
            eid = child.get(qn('w:id'))
            etag = TagModel(type="endnote", xml_attributes={"id": eid})
            tid = add_tag(etag)
            full_text += f"<{tid}></{tid}>"

        # 7. Deleted Text (w:del)
        elif tag_name == qn('w:del'):
            continue
            
    # Flush remaining
    flush_run_buffer()

    if not full_text:
        return []

    # Decision: Smart Splitting
    # Always split sentences, then repair tags across boundaries.
    
    segments_to_create = []
    final_segments = []

    # 1. Split (using SemanticAligner which protects tags from breaking splitting logic)
    sentences = _split_sentences(full_text, context.get("segmentation_func"), lang=context.get("source_lang", "en"))
    
    # 2. Repair Tags (Clone open tags across split boundaries)
    repaired_sentences = _repair_tags(sentences)
    
    for idx, sentence in enumerate(repaired_sentences):
         segments_to_create.append((sentence, idx))

    for content, sub_index in segments_to_create:
        # Create unique location
        seg_loc = location.copy()
        seg_loc["sub_index"] = sub_index

        # Detect Leading/Trailing Spaces
        src_leading = ""
        src_trailing = ""
        
        # Leading
        m_lead = re.match(r'^(\s+)', content)
        if m_lead:
            src_leading = m_lead.group(1)
            
        # Trailing
        m_trail = re.search(r'(\s+)$', content)
        if m_trail:
            src_trailing = m_trail.group(1)
            
        if src_leading or src_trailing:
            seg_loc["whitespaces"] = {
                "leading": src_leading,
                "trailing": src_trailing
            }
        
        final_segments.append(SegmentInternal(
            id=str(uuid.uuid4()),
            segment_id=str(uuid.uuid4()),
            source_text=content,
            target_content=None, # Translation starts empty/null
            status="draft",
            tags=tags, # Pass tags to all segments so split parts can reference them
            metadata=seg_loc
        ))

    return final_segments

import pysbd

_segmenter_cache = {}

def _get_segmenter(lang: str):
    if lang not in _segmenter_cache:
        # Pysbd supports ISO codes. Fallback to en.
        try:
             _segmenter_cache[lang] = pysbd.Segmenter(language=lang, clean=False)
        except:
             _segmenter_cache[lang] = pysbd.Segmenter(language="en", clean=False)
    return _segmenter_cache[lang]

def _split_sentences(text: str, segmentation_func=None, lang="en") -> List[str]:
    # Use pysbd for robust splitting
    if segmentation_func:
        return segmentation_func(text)
    
    seg = _get_segmenter(lang)
    return seg.segment(text)

def _extract_tags(run_element) -> List[TagModel]:
    """
    Inspects a run XML ELEMENT for formatting.
    Returns a list of detected TagModels.
    """
    found = []
    
    # Needs valid namespaces
    namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    rPr = run_element.find(qn('w:rPr'))
    if rPr is None:
        return found
        
    # Bold
    # w:b can be present (True) or have w:val="0"/"false" (False)
    # Simple check: availability
    if rPr.find(qn('w:b')) is not None:
         # To be pedantic: check w:val
         # val = rPr.find(qn('w:b')).get(qn('w:val'))
         # if val not in ['0', 'false', 'off']:
         found.append(TagModel(type="bold"))
         
    if rPr.find(qn('w:i')) is not None:
        found.append(TagModel(type="italic"))
        
    if rPr.find(qn('w:u')) is not None:
        found.append(TagModel(type="underline"))
        
    # Superscript / Subscript (w:vertAlign w:val='superscript')
    vAlign = rPr.find(qn('w:vertAlign'))
    if vAlign is not None:
        val = vAlign.get(qn('w:val'))
        if val == 'superscript':
            found.append(TagModel(type="superscript"))
        elif val == 'subscript':
            found.append(TagModel(type="subscript"))
            
    # Color
    color = rPr.find(qn('w:color'))
    if color is not None:
        val = color.get(qn('w:val'))
        # Ignore auto or black (common defaults that cause noise)
        if val and val != 'auto' and val != '000000':
            found.append(TagModel(type="color", xml_attributes={"color": val}))

    # Highlighting
    highlight = rPr.find(qn('w:highlight'))
    if highlight is not None:
        val = highlight.get(qn('w:val'))
        found.append(TagModel(type="highlight", xml_attributes={"color": str(val)}))

    # --- NEW: Extended Font Properties ---
    
    # Font Size (w:sz) - Value is in half-points (e.g., 24 = 12pt)
    sz = rPr.find(qn('w:sz'))
    if sz is not None:
        val = sz.get(qn('w:val'))
        if val:
            found.append(TagModel(type="size", xml_attributes={"val": str(val)}))
            
    # Font Family (w:rFonts)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is not None:
        # Prefer ascii -> hAnsi -> eastAsia -> cs
        font_name = rFonts.get(qn('w:ascii')) or rFonts.get(qn('w:hAnsi')) or rFonts.get(qn('w:eastAsia'))
        if font_name:
             found.append(TagModel(type="font", xml_attributes={"name": str(font_name)}))

    # Strikethrough (w:strike)
    strike = rPr.find(qn('w:strike'))
    if strike is not None:
        # Check explicit val (default is true if missing)
        val = strike.get(qn('w:val'))
        if val not in ['0', 'false', 'off']:
            found.append(TagModel(type="strike"))
            
    # Small Caps (w:smallCaps)
    smallCaps = rPr.find(qn('w:smallCaps'))
    if smallCaps is not None:
        val = smallCaps.get(qn('w:val'))
        if val not in ['0', 'false', 'off']:
            found.append(TagModel(type="smallCaps"))

    # All Caps (w:caps)
    caps = rPr.find(qn('w:caps'))
    if caps is not None:
        val = caps.get(qn('w:val'))
        if val not in ['0', 'false', 'off']:
            found.append(TagModel(type="caps"))

    return found

def _get_run_text(run_element, namespaces):
    """
    Extracts text from a run, handling <w:t>, <w:br>, <w:tab>.
    """
    text = ""
    for child in run_element:
        tag = child.tag
        if tag == qn('w:t'):
             text += child.text or ""
        elif tag == qn('w:br'):
             text += "\n" # Or special placeholder? Usually break is fine as newline
        elif tag == qn('w:tab'):
             text += "\t"
        elif tag == qn('w:cr'):
             text += "\n"
    return text

def _get_run_signature(run_element):
    """
    Returns a hashable signature of meaningful formatting properties.
    Ignores rsid, lang, etc.
    """
    rPr = run_element.find(qn('w:rPr'))
    if rPr is None:
        return None # Default style
        
    # We essentially need to serialize the properties we CARE about.
    # The list matches _extract_tags logic.
    
    sig = []
    
    # 1. Boolean Toggles
    for tag in ['b', 'i', 'u', 'strike', 'smallCaps', 'caps', 'vanish', 'webHidden']:
        el = rPr.find(qn(f'w:{tag}'))
        if el is not None:
            val = el.get(qn('w:val'))
            # 'on', '1', 'true', or missing attribute = True
            # 'off', '0', 'false' = False
            state = True
            if val in ['0', 'false', 'off']: state = False
            sig.append((tag, state))
    
    # 2. Valued Properties
    # (tag_name, attr_name)
    valued_props = [
        ('color', 'val'),
        ('highlight', 'val'),
        ('sz', 'val'),
        ('rFonts', 'ascii'), # Simplified: track ascii font as proxy
        ('rFonts', 'hAnsi'),
        ('vertAlign', 'val'),
        ('shd', 'fill'), # Shading
        ('kern', 'val'),
        ('position', 'val'),
    ]
    
    for tag, attr in valued_props:
        el = rPr.find(qn(f'w:{tag}'))
        if el is not None:
            val = el.get(qn(f'w:{attr}'))
            if val:
                sig.append((tag, attr, val))
                
    return tuple(sorted(sig))

def _is_pure_text_run(run_element):
    """
    Returns True if the run contains only text-like elements (t, tab, br, cr, rPr).
    Returns False if it contains drawings, footnotes, etc.
    """
    allowed = [qn('w:t'), qn('w:br'), qn('w:tab'), qn('w:cr'), qn('w:rPr')]
    for child in run_element:
        if child.tag not in allowed:
            return False
    return True

def _handle_shape_element(shape_element, add_tag_func, context) -> str:
    # NEW: Drill down for Textboxes (w:txbxContent)
    found_textbox = False
    shape_id = None
    
    # Check for textbox content
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    txbx_contents = shape_element.findall('.//w:txbxContent', ns)
    
    if txbx_contents:
        logger.debug(f"Shape {shape_element}: Found {len(txbx_contents)} textboxes")
        shape_id = str(uuid.uuid4())
        found_textbox = True
        
        for txbx in txbx_contents:
            # Iterate paragraphs
            for i, para in enumerate(txbx.findall('.//w:p', ns)):
                loc = {
                    "type": "shape",
                    "shape_id": shape_id,
                    "p_index": i
                }
                
                sub_segments = _process_paragraph(para, loc, context)
                if sub_segments:
                    context["extra_segments"].extend(sub_segments)

    # Create shape tag
    if found_textbox:
        shape_tag = TagModel(type="shape", content="[SHAPE]", xml_attributes={"id": shape_id})
    else:
        shape_tag = TagModel(type="shape", content="[SHAPE]")
        
    tid = add_tag_func(shape_tag)
    return f"<{tid}></{tid}>"


def _process_run_element(run_element, add_tag_func, context, text_override=None) -> str:
    """
    Helper to process a w:r element.
    Handles formatting (Bold/Italic), Embedded Comments, Tabs, Breaks, and Text.
    Also handles URL regex detection.
    """
    
    # 1. Helper for checking URL in text
    def process_text_for_urls(txt):
        if not txt: return ""
        # Improved regex to avoid catastrophic backtracking or false positives
        url_pattern = re.compile(r'(https?://[^\s<>"]+|www\.[^\s<>"]+)')
        parts = url_pattern.split(txt)
        res = ""
        if len(parts) > 1:
            for part in parts:
                if url_pattern.match(part):
                        l_tag = TagModel(type="link", xml_attributes={"is_hyperlink": True})
                        l_tid = add_tag_func(l_tag)
                        res += f"<{l_tid}>{part}</{l_tid}>"
                elif part:
                        res += part
            return res
        else:
            return txt

    # 2. Extract formatting from Run XML directly
    extracted_tags = _extract_tags(run_element)
    
    active_ids = []
    full_run_text = ""
    
    # Open formatting tags
    if extracted_tags:
        for t in extracted_tags:
            tid = add_tag_func(t)
            full_run_text += f"<{tid}>"
            active_ids.append(tid)
            
    # 3. Process Content
    run_content = ""
    
    if text_override is not None:
        # Override mode: Caller merged text for us.
        run_content = process_text_for_urls(text_override)
    else:
        # Standard mode: Iterate children
        for child in run_element:
            tag_name = child.tag
            
            if tag_name == qn('w:t'):
                run_content += process_text_for_urls(child.text or "")
                
            elif tag_name == qn('w:tab'):
                 run_content += "\t"
                 
            elif tag_name == qn('w:br') or tag_name == qn('w:cr'):
                 run_content += "\n"
                 
            elif tag_name == qn('w:commentReference'):
                cid = child.get(qn('w:id'))
                was_handled = "_handled_ranges" in context and cid in context["_handled_ranges"]
                is_active = "_active_ranges" in context and cid in context["_active_ranges"]
                
                if not was_handled and not is_active and cid and context["comments_map"].get(cid):
                    ctext = context["comments_map"][cid]
                    com_tag = TagModel(type="comment", content=ctext, ref_id=cid)
                    tid = add_tag_func(com_tag)
                    run_content += f"<{tid}></{tid}>"

            elif tag_name == qn('w:drawing') or tag_name == qn('w:pict'):
                 run_content += _handle_shape_element(child, add_tag_func, context)

            elif tag_name == "{http://schemas.openxmlformats.org/markup-compatibility/2006}AlternateContent":
                 mc_ns = {'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006', 'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                 choice = child.find('mc:Choice', mc_ns)
                 fallback = child.find('mc:Fallback', mc_ns)
                 target_el = None
                 if choice is not None:
                     target_el = choice.find('.//w:drawing', mc_ns)
                 if target_el is None and fallback is not None:
                     target_el = fallback.find('.//w:pict', mc_ns)
                 
                 if target_el is not None:
                     run_content += _handle_shape_element(target_el, add_tag_func, context)

            elif tag_name == qn('w:footnoteReference'):
                fid = child.get(qn('w:id'))
                ftag = TagModel(type="footnote", xml_attributes={"id": fid})
                tid = add_tag_func(ftag)
                run_content += f"<{tid}></{tid}>"

            elif tag_name == qn('w:endnoteReference'):
                eid = child.get(qn('w:id'))
                etag = TagModel(type="endnote", xml_attributes={"id": eid})
                tid = add_tag_func(etag)
                run_content += f"<{tid}></{tid}>"

    full_run_text += run_content
    
    # Close formatting tags (Reverse order)
    for tid in reversed(active_ids):
        full_run_text += f"</{tid}>"
        
    return full_run_text

def _repair_tags(segments: List[str]) -> List[str]:
    """
    Ensures that if a segment ends with open tags, they are closed,
    and reopened in the next segment.
    """
    repaired = []
    stack = []
    # Regex to find tags: <1>, </1>
    pattern = re.compile(r'<(/?(\d+))>')
    
    for part in segments:
        # 1. Prepend Open Tags from Stack (Re-Open)
        prefix = "".join([f"<{tid}>" for tid in stack])
        current_seg = prefix + part
        
        # 2. Update Stack based on tags in THIS part (Original content)
        # We must scan 'part' to avoiding seeing the tags we just prepended
        for m in pattern.finditer(part):
            full_tag = m.group(1) # "1" or "/1"
            tid = m.group(2)
            is_close = full_tag.startswith("/")
            
            if is_close:
                # Attempt to pop from stack
                if stack and stack[-1] == tid:
                    stack.pop()
            else:
                stack.append(tid)
        
        # 3. Append Close Tags for remaining Stack (Close)
        suffix = "".join([f"</{tid}>" for tid in reversed(stack)])
        current_seg = current_seg + suffix
        
        repaired.append(current_seg)
        
    return repaired
